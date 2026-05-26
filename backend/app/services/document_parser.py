import platform
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader

ALLOWED_RESUME_EXTENSIONS = {".pdf", ".doc", ".docx"}
MAX_RESUME_BYTES = 10 * 1024 * 1024  # 10 MB


def validate_resume_file(filename: str | None, size: int) -> str:
    if not filename:
        raise HTTPException(status_code=400, detail="Filename is required")
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_RESUME_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed: {', '.join(sorted(ALLOWED_RESUME_EXTENSIONS))}",
        )
    if size > MAX_RESUME_BYTES:
        raise HTTPException(status_code=400, detail="File too large (max 10 MB)")
    return ext


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text.strip())
    return "\n\n".join(parts)


def extract_text_from_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_doc_with_textutil(data: bytes) -> str:
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = Path(tmpdir) / "resume.doc"
        doc_path.write_bytes(data)
        out_base = Path(tmpdir) / "resume"
        subprocess.run(
            ["textutil", "-convert", "txt", "-output", str(out_base), str(doc_path)],
            check=True,
            capture_output=True,
            timeout=30,
        )
        txt_path = Path(f"{out_base}.txt")
        if not txt_path.exists():
            raise RuntimeError("textutil did not produce output")
        return txt_path.read_text(encoding="utf-8", errors="replace")


def _extract_doc_with_antiword(data: bytes) -> str:
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = Path(tmpdir) / "resume.doc"
        doc_path.write_bytes(data)
        result = subprocess.run(
            ["antiword", str(doc_path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=30,
        )
        return result.stdout


def extract_text_from_doc(data: bytes) -> str:
    if platform.system() == "Darwin":
        try:
            return _extract_doc_with_textutil(data)
        except (subprocess.CalledProcessError, OSError, RuntimeError):
            pass

    try:
        return _extract_doc_with_antiword(data)
    except FileNotFoundError as exc:
        raise HTTPException(
            status_code=400,
            detail=(
                "Legacy .doc files require macOS textutil or Linux antiword. "
                "Please save as .docx or .pdf and upload again."
            ),
        ) from exc
    except subprocess.CalledProcessError as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse .doc file: {exc}",
        ) from exc


def extract_text_from_document(data: bytes, filename: str) -> str:
    ext = validate_resume_file(filename, len(data))

    if ext == ".pdf":
        try:
            content = extract_text_from_pdf(data)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {exc}") from exc
    elif ext == ".docx":
        try:
            content = extract_text_from_docx(data)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {exc}") from exc
    else:
        content = extract_text_from_doc(data)

    if len(content.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="Resume content too short or unreadable. Try a different file format.",
        )

    return content.strip()
